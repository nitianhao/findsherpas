"""
sales_materials_generator.py

Generates three sales derivative documents from a completed AuditReport,
saved alongside the main report files in the same output directory.

Outputs:
  {slug}_exec_summary.docx  — 2-page Executive Summary (for prospects)
  {slug}_brief.docx         — 1-page Forwardable Brief (for VP/CMO)
  {slug}_cold_email.txt     — Cold email opener, 3 versions (plain text)

Designed to be generic: works for any site, any capability set,
any mix of failure modes. All examples are drawn from live audit data.
"""

from __future__ import annotations

import logging
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv

from src.models import AuditReport, Severity

load_dotenv(override=True)

logger = logging.getLogger(__name__)

_MODEL = "claude-sonnet-4-20250514"

# ── Brand colours ────────────────────────────────────────────────────────────
_TEAL = RGBColor(0x00, 0x9B, 0xA3)   # Find Sherpas primary
_DARK = RGBColor(0x0A, 0x0A, 0x0A)   # Near-black body text
_MUTED = RGBColor(0x73, 0x73, 0x73)  # Muted grey for meta text

# ── Internal severity ordering ───────────────────────────────────────────────
_SEVERITY_RANK: dict[str, int] = {
    Severity.CRITICAL.value: 3,
    Severity.MODERATE.value: 2,
    Severity.MINOR.value:    1,
    Severity.PASS.value:     0,
}

_SEVERITY_SHORT: dict[str, str] = {
    Severity.CRITICAL.value: "Critical",
    Severity.MODERATE.value: "Moderate",
    Severity.MINOR.value:    "Minor",
    Severity.PASS.value:     "Pass",
}

_CAPABILITY_NAMES: dict[str, str] = {
    "TYPO_TOLERANCE":       "Typo Tolerance",
    "LANGUAGE_UNDERSTANDING": "Language Understanding",
    "PRODUCT_DISCOVERY":    "Product Discovery",
    "BRAND_MODEL_SEARCH":   "Brand & Model Search",
    "FILTERS_CONSTRAINTS":  "Filters & Constraints",
    "SHOPPING_CONTEXT":     "Shopping Context",
}


# ─────────────────────────────────────────────────────────────────────────────
# 1. DATA EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def _extract_sales_context(report: AuditReport) -> dict:
    """
    Compute the key stats and pick the most compelling findings from an
    AuditReport.  Returns a plain dict that feeds directly into the prompt
    templates — no report-specific logic lives in the prompts themselves.
    """
    all_judgments = [j for cs in report.capability_scores for j in cs.judgments]
    total_queries = len(all_judgments)

    if total_queries == 0:
        raise ValueError("AuditReport contains no judgments.")

    # ── Capability pass count ────────────────────────────────────────────────
    capabilities_passed = sum(
        1 for cs in report.capability_scores
        if cs.severity == Severity.PASS.value
    )
    total_capabilities = len(report.capability_scores)

    # ── Position stats ───────────────────────────────────────────────────────
    avg_best_position = round(
        sum(j.displacement + 1 for j in all_judgments) / total_queries, 1
    )
    outside_top3_count = sum(1 for j in all_judgments if j.displacement > 2)
    pct_outside_top3 = round(outside_top3_count / total_queries * 100)
    pct_in_top3 = 100 - pct_outside_top3

    # ── Failure mode breakdown ───────────────────────────────────────────────
    ranking_failures = sum(
        1 for j in all_judgments if j.failure_mode == "POOR_RANKING"
    )
    pct_ranking_failure = round(ranking_failures / total_queries * 100)

    # ── Critical capabilities ────────────────────────────────────────────────
    critical_caps = [
        _CAPABILITY_NAMES.get(cs.capability, cs.capability)
        for cs in report.capability_scores
        if cs.severity == Severity.CRITICAL.value
    ]

    # ── Top findings: non-pass, sorted worst-first ───────────────────────────
    # Primary sort: severity (Critical > Moderate > Minor)
    # Secondary sort: effective displacement — NO_FUZZY_MATCHING and
    # ZERO_RESULTS_OR_GARBAGE get displacement=9999 because their
    # actual displacement is 0 (garbage result sits at #1) but the
    # failure is more egregious than a ranking miss.
    _EGREGIOUS_MODES = {"NO_FUZZY_MATCHING", "ZERO_RESULTS_OR_GARBAGE"}

    def _effective_displacement(j) -> int:
        return 9999 if j.failure_mode in _EGREGIOUS_MODES else j.displacement

    non_pass = [j for j in all_judgments if j.severity != Severity.PASS.value]
    sorted_judgments = sorted(
        non_pass,
        key=lambda j: (-_SEVERITY_RANK.get(j.severity, 0), -_effective_displacement(j)),
    )

    top_findings = []
    for j in sorted_judgments[:4]:
        # What the customer actually saw (top 3 by original rank)
        by_rank = sorted(j.results, key=lambda r: r.original_rank)
        customer_saw = []
        for r in by_rank[:3]:
            price_str = f" ({r.price})" if r.price else ""
            customer_saw.append(f"#{r.original_rank} {r.title}{price_str}")

        # Best available match (highest relevance score)
        best = max(j.results, key=lambda r: r.relevance_score) if j.results else None
        best_str = f'"{best.title}"' if best else "unknown"
        if best and best.price:
            best_str += f" ({best.price})"

        top_findings.append({
            "query":        j.test_query.query,
            "severity":     _SEVERITY_SHORT.get(j.severity, j.severity),
            "best_position": j.displacement + 1,
            "displacement": j.displacement,
            "failure_mode": j.failure_mode,
            "evidence":     j.evidence,
            "customer_saw": customer_saw,
            "best_match":   best_str,
        })

    return {
        "site_name":           report.site_context.site_name or report.site_context.url,
        "site_url":            report.site_context.url,
        "capabilities_passed": capabilities_passed,
        "total_capabilities":  total_capabilities,
        "total_queries":       total_queries,
        "avg_best_position":   avg_best_position,
        "pct_outside_top3":    pct_outside_top3,
        "pct_in_top3":         pct_in_top3,
        "ranking_failures":    ranking_failures,
        "pct_ranking_failure": pct_ranking_failure,
        "critical_caps":       critical_caps,
        "top_findings":        top_findings,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 2. PROMPT BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def _fmt_finding(f: dict) -> str:
    """Format a single finding dict as a readable block for prompt injection."""
    saw = "\n".join(f"    - {r}" for r in f["customer_saw"])
    return (
        f'  Query: "{f["query"]}"\n'
        f'  Severity: {f["severity"]}\n'
        f'  Best result position: #{f["best_position"]} '
        f'(displaced {f["displacement"]} positions from #1)\n'
        f'  Failure mode: {f["failure_mode"]}\n'
        f'  Evidence: {f["evidence"]}\n'
        f'  What customer saw (top 3 results):\n{saw}\n'
        f'  Best available match: {f["best_match"]}'
    )


def _exec_summary_prompt(ctx: dict) -> str:
    findings_text = "\n\n".join(
        f"Finding {i + 1}:\n{_fmt_finding(f)}"
        for i, f in enumerate(ctx["top_findings"])
    )
    critical_str = (
        ", ".join(ctx["critical_caps"]) if ctx["critical_caps"]
        else "multiple capabilities"
    )

    return f"""You are writing a 2-page Executive Summary for Find Sherpas (findsherpas.com), \
a boutique search optimization agency. This document goes to a prospect at {ctx["site_name"]} \
after they reply to a cold email.

BRAND VOICE: Sharp, peer-to-peer, consultative. Short sentences. Direct assertions. \
No hedging. No "leverage," "synergy," "optimize," "robust," "exciting," or "deep dive." \
No methodology section. No filler. Let the data speak — avoid obvious editorial comments \
like "this is a problem for a pharmacy" or similar.

OUTPUT FORMAT: Use ## for section headings, ### for named sub-findings. \
Prose paragraphs only — no bullet point lists anywhere in the body.

REQUIRED SECTIONS (use these exact headings):

## What We Found
3–4 sentences. Open with the pass rate ("{ctx["capabilities_passed"]} of \
{ctx["total_capabilities"]}"). Name the average best-result position \
(#{ctx["avg_best_position"]}) and the top-3 relevance rate ({ctx["pct_in_top3"]}%). \
Close with the fixability framing: this is a ranking problem, not a catalog problem.

## The Findings That Matter
3–4 named sub-findings (use ### for each). Each must use a SPECIFIC example \
from the audit data below — name the query, name what came back, name the position. \
2–4 sentences per finding. At least one finding must include the exact best-result \
position number.

## The Revenue Picture
3–4 sentences. Use this logic: search converts at 1.8–3× vs browse-only visitors; \
top 3 captures ~60% of clicks; the gap between {ctx["pct_in_top3"]}% and the 80%+ \
industry target is where the revenue is. Keep it directional — no fabricated euro amounts. \
Include the fixability point: ranking configuration, no catalog rebuild, no front-end work. \
Comparable sites recover 15–25% of lost search CVR once ranking is fixed.

## Next Step
2 sentences max. Soft CTA — offer 30 minutes to walk through the top 3 prioritized fixes \
ranked by effort-to-impact. Frame it as the prospect's decision, not a pitch.

--- AUDIT DATA (use these findings — do not invent examples) ---

Site: {ctx["site_name"]} ({ctx["site_url"]})

Key stats:
  Capabilities passed: {ctx["capabilities_passed"]} of {ctx["total_capabilities"]}
  Total queries tested: {ctx["total_queries"]}
  Average best-result position: #{ctx["avg_best_position"]} (industry target: #1–2)
  Relevant result in top 3: {ctx["pct_in_top3"]}% (industry target: 80%+)
  Best result outside top 3: {ctx["pct_outside_top3"]}% of queries
  Ranking failures (retrieved correctly, ranked wrong): \
{ctx["ranking_failures"]} of {ctx["total_queries"]} ({ctx["pct_ranking_failure"]}%)
  Critical capabilities: {critical_str}

Top findings (sorted by severity then displacement):
{findings_text}

Industry benchmarks:
  Search users convert at 1.8–3× the rate of browse-only visitors (Econsultancy, Algolia)
  Top 3 results capture ~55–60% of all search clicks
  80%+ relevant result in top 3 is the industry target (Baymard Institute)
  Fixing the ranking layer typically recovers 15–25% of lost search CVR

Write the full document now.
"""


def _brief_prompt(ctx: dict) -> str:
    top = ctx["top_findings"][0] if ctx["top_findings"] else {}
    saw_str = ", ".join(top.get("customer_saw", [])[:3])

    return f"""You are writing a 1-page internal memo for Find Sherpas (findsherpas.com). \
This document is designed to be forwarded by a mid-level contact to their VP or CMO \
at {ctx["site_name"]}. It must fit on one page and be readable in 90 seconds. \
The reader has not seen the full audit.

OUTPUT FORMAT: Strict memo format, then four short paragraphs. \
Plain prose — no bullet points, no headers beyond the memo block.

Start with this exact header block:
TO: [leave blank]
FROM: [Name], Find Sherpas
RE: {ctx["site_name"]} — Search Revenue Exposure
DATE: [current month and year]

Then write exactly four paragraphs:

Paragraph 1 (3 sentences): Open with the headline stat — \
{ctx["pct_in_top3"]}% relevant result in top 3 vs the 80%+ industry benchmark. \
State the mechanism: search visitors convert at 1.8–3× vs browse-only visitors, \
and positions 1–3 capture ~60% of clicks. Close with the consequence: \
the majority of highest-intent traffic is being directed to the wrong products.

Paragraph 2 (3–4 sentences): One concrete example from the finding below. \
Name the exact query. Name what came back. Include the best-result position number. \
Close with a single punchy consequence line \
(e.g. "That's a buying customer turned into a scrolling customer").

Paragraph 3 (2–3 sentences): State that this pattern repeats — \
{ctx["ranking_failures"]} of {ctx["total_queries"]} queries \
({ctx["pct_ranking_failure"]}%) followed the same pattern. \
State the fix: ranking configuration, not a platform change, not a catalog rebuild.

Paragraph 4 (2 sentences): The contact has the full audit. \
Offer 30 minutes to walk through the top 3 fixes ranked by effort-to-impact ratio.

TONE: Executive peer-to-peer. Short and cold. No agency-speak. \
No "we help companies." Assume the reader is smart and busy.

--- FINDING TO USE ---
Query: "{top.get("query", "")}"
Severity: {top.get("severity", "")}
Best result position: #{top.get("best_position", "")}
What customer saw (top 3): {saw_str}
Evidence: {top.get("evidence", "")}

Write the full memo now.
"""


def _cold_email_prompt(ctx: dict) -> str:
    findings = ctx["top_findings"]
    top    = findings[0] if len(findings) > 0 else {}
    second = findings[1] if len(findings) > 1 else {}

    top_saw    = ", ".join(top.get("customer_saw", [])[:3])
    second_saw = ", ".join(second.get("customer_saw", [])[:3])

    return f"""You are writing the opening 2–4 sentences of a cold outreach email \
for Find Sherpas (findsherpas.com), a boutique search optimization agency. \
The recipient works at {ctx["site_name"]}. \
This snippet is the first thing they read — before any pitch, any intro, any credentials. \
Its job is to earn 10 more seconds of attention by naming something specific and \
surprising about their own site.

RULES:
- Start mid-thought. No "I came across your site," no "I wanted to reach out," \
no "hope this finds you well."
- Name a specific, verifiable finding — something the reader can check right now.
- Hint at a revenue or conversion consequence without overstating it.
- Must not read like it could be sent to any other company.
- No fluff words: no "fascinating," "impressive," "excited," "unique," "powerful."
- Do not explain who Find Sherpas is — that comes later in the email.
- Plain text only. 2–4 sentences per version. No subject line. No sign-off.

Write exactly 3 versions, labelled Version A, Version B, Version C.

Version A — use this finding:
  Query: "{top.get("query", "")}"
  What customer saw (top 3): {top_saw}
  Best result was at position: #{top.get("best_position", "")}
  Evidence: {top.get("evidence", "")}

Version B — use this finding:
  Query: "{second.get("query", "")}"
  What customer saw (top 3): {second_saw}
  Best result was at position: #{second.get("best_position", "")}
  Evidence: {second.get("evidence", "")}

Version C — use the aggregate picture:
  Relevant result in top 3: {ctx["pct_in_top3"]}% across {ctx["total_queries"]} \
queries (industry target: 80%+)
  Average best-result position: #{ctx["avg_best_position"]}
  {ctx["ranking_failures"]} of {ctx["total_queries"]} queries: right product \
retrieved, ranked wrong.

Write all three versions now.
"""


# ─────────────────────────────────────────────────────────────────────────────
# 3. CLAUDE API CALL
# ─────────────────────────────────────────────────────────────────────────────

def _call_claude(prompt: str) -> str:
    client = anthropic.Anthropic()
    response = client.messages.create(
        model=_MODEL,
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# 4. DOCX WRITER
# ─────────────────────────────────────────────────────────────────────────────

def _set_font(
    run,
    size_pt: float,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_header_rule(doc: Document, site_name: str, doc_type: str) -> None:
    """Branded header: agency name + document type + site name."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("Find Sherpas")
    _set_font(r1, 10, bold=True, color=_TEAL)
    r2 = p.add_run(f"  |  {doc_type}  |  {site_name}")
    _set_font(r2, 10, color=_MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    rr = rule.add_run("─" * 80)
    _set_font(rr, 7, color=_TEAL)


def _add_footer(doc: Document) -> None:
    """Centred URL footer."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("findsherpas.com")
    _set_font(r, 9, color=_MUTED)


def _write_docx(content: str, site_name: str, doc_type: str, out_path: Path) -> None:
    """
    Convert Claude's markdown-structured output to a clean .docx.

    Parsing rules:
      ## heading   → section heading (teal, bold, 12pt)
      ### heading  → sub-heading (dark, bold, 11pt)
      TO:/FROM:/RE:/DATE: lines → memo header (bold label + normal value)
      blank line   → paragraph separator
      everything else → body paragraph (10pt)
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default paragraph spacing
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    _add_header_rule(doc, site_name, doc_type)

    memo_header_keys = ("TO:", "FROM:", "RE:", "DATE:")

    for line in content.split("\n"):
        stripped = line.strip()

        # ── Blank lines become breathing room (don't add empty paragraphs)
        if not stripped:
            continue

        # ── ## Section heading
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            _set_font(r, 12, bold=True, color=_TEAL)

        # ── ### Sub-heading (named findings)
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(text)
            _set_font(r, 10, bold=True, color=_DARK)

        # ── Memo header lines (TO: / FROM: / RE: / DATE:)
        elif any(stripped.startswith(k) for k in memo_header_keys):
            label, _, rest = stripped.partition(":")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_label = p.add_run(label + ":")
            _set_font(r_label, 10, bold=True, color=_DARK)
            r_rest = p.add_run(rest)
            _set_font(r_rest, 10, color=_DARK)

        # ── Body paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            r = p.add_run(stripped)
            _set_font(r, 10, color=_DARK)

    _add_footer(doc)
    doc.save(str(out_path))
    logger.info("  Saved: %s", out_path)


# ─────────────────────────────────────────────────────────────────────────────
# 5. PUBLIC ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def generate_sales_materials(
    report: AuditReport,
    out_dir: Path,
    slug: str,
) -> None:
    """
    Generate the three sales derivative documents from a completed AuditReport.

    Args:
        report:  Completed AuditReport (the same object returned by generate_report()).
        out_dir: Output directory — same folder as the main report files.
        slug:    File slug, e.g. "shop_apotheke_com_20260410_100841".
                 Used as the filename prefix for all three output files.
    """
    logger.info("Generating sales materials for %s...", report.site_context.site_name)

    try:
        ctx = _extract_sales_context(report)
    except ValueError as exc:
        logger.warning("Sales materials skipped: %s", exc)
        return

    site_name = ctx["site_name"]

    # ── Executive Summary (.docx) ─────────────────────────────────────────
    logger.info("  [1/3] Executive Summary...")
    exec_text = _call_claude(_exec_summary_prompt(ctx))
    _write_docx(
        exec_text,
        site_name=site_name,
        doc_type="Executive Summary",
        out_path=out_dir / f"{slug}_exec_summary.docx",
    )

    # ── Forwardable Brief (.docx) ─────────────────────────────────────────
    logger.info("  [2/3] Forwardable Brief...")
    brief_text = _call_claude(_brief_prompt(ctx))
    _write_docx(
        brief_text,
        site_name=site_name,
        doc_type="Forwardable Brief",
        out_path=out_dir / f"{slug}_brief.docx",
    )

    # ── Cold Email Snippet (.txt) ─────────────────────────────────────────
    logger.info("  [3/3] Cold email snippet...")
    cold_text = _call_claude(_cold_email_prompt(ctx))
    cold_path = out_dir / f"{slug}_cold_email.txt"
    cold_path.write_text(cold_text, encoding="utf-8")
    logger.info("  Saved: %s", cold_path)

    logger.info("Sales materials complete — 3 files written to %s", out_dir)
